from pdfminer.high_level import extract_text
from docx import Document
from typing import cast
import pandas as pd
import io
import re
import logging
import os

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def parse_file(file_path_or_object, filename=None):
    """Parse uploaded file and extract text chunks"""
    try:
        if isinstance(file_path_or_object, str):
            # It's a file path
            if filename is None:
                filename = os.path.basename(file_path_or_object).lower()
            with open(file_path_or_object, 'rb') as f:
                file_content = f.read()
        else:
            # It's a FastAPI UploadFile object
            filename = file_path_or_object.filename.lower()
            file_content = file_path_or_object.file.read()
        
        if filename.endswith('.pdf'):
            return parse_pdf(file_content)
        elif filename.endswith('.docx'):
            return parse_docx(file_content)
        elif filename.endswith('.csv'):
            return parse_csv(file_content)
        elif filename.endswith('.txt'):
            return parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
            
    except Exception as e:
        logger.error(f"Error parsing file {filename}: {e}")
        raise

def parse_pdf(file_content):
    """Parse PDF file"""
    try:
        # Use BytesIO to create a file-like object
        pdf_file = io.BytesIO(file_content)
        text = extract_text(pdf_file)
        return chunk_text(text)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return []

def parse_docx(file_content):
    """Parse DOCX file"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_parts)
        return chunk_text(full_text)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return []

def parse_csv(file_content):
    """Parse CSV file"""
    try:
        csv_file = io.BytesIO(file_content)
        
        df = None
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
            try:
                csv_file.seek(0)
                df = pd.read_csv(csv_file, encoding=encoding, on_bad_lines='skip')
                logger.info(f"Successfully parsed CSV with {encoding} encoding")
                break
            except (UnicodeDecodeError, pd.errors.ParserError) as e:
                logger.debug(f"Failed {encoding} encoding: {e}")
                continue
        
        if df is None or df.empty:
            raise ValueError("Unable to parse CSV with any known encoding or file is empty")
        
        # Convert to text chunks
        chunks = []
        
        # Add header information
        headers = ', '.join(df.columns.tolist())
        chunks.append(f"CSV Headers: {headers}")
        
        # Add rows as chunks (limit to avoid too many chunks)
        max_rows = 100  # Limit rows to process
        for i, row in df.head(max_rows).iterrows():
            row_idx = int(cast(int, i)) + 1

            row_text = ', '.join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
            chunks.append(f"Row {row_idx}: {row_text}")
            
            # If row is too long, split it
            if len(row_text) > 500:
                 for subchunk in chunk_text(row_text, 200):
                   chunks.append(subchunk)
        
        if len(df) > max_rows:
            chunks.append(f"... and {len(df) - max_rows} more rows")
        
        return chunks
    except Exception as e:
        logger.error(f"Error parsing CSV: {e}")
        return []

def parse_txt(file_content):
    """Parse text file"""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'ascii']:
            try:
                text = file_content.decode(encoding)
                return chunk_text(text)
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode text file with common encodings")
    except Exception as e:
        logger.error(f"Error parsing text file: {e}")
        return []


def chunk_text(text, chunk_size=800, overlap=100):
    """Split text into overlapping chunks for better context"""
    if not text or not text.strip():
        return []
    
    # Better cleaning - remove multiple newlines and extra spaces
    text = re.sub(r'\n+', '\n', text)  # Replace multiple newlines with single
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Remove common PDF artifacts
    text = re.sub(r'\x0c', '', text)  # Form feed character
    text = re.sub(r'\.{2,}', '...', text)  # Multiple dots
    
    # If text is shorter than chunk_size, return as single chunk
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    sentences = re.split(r'(?<=[.!?])\s+', text)  # Split on sentences
    
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += " " + sentence if current_chunk else sentence
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    # Ensure overlap between chunks
    if len(chunks) > 1:
        overlapped_chunks = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i-1]
            overlap_text = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
            current_chunk = chunks[i]
            overlapped_chunks.append(overlap_text + " " + current_chunk)
        return overlapped_chunks
    
    return chunks